"""生成项目详解 Word 文档 —— 面向初学者的完整指南"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ==================== 页面设置 ====================
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

def add_para(text, bold=False, size=11, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.8)
    return p

def add_term(term, explanation):
    """添加术语解释"""
    p = doc.add_paragraph()
    run_term = p.add_run(f"■ {term}：")
    run_term.bold = True
    run_term.font.size = Pt(11)
    run_expl = p.add_run(explanation)
    run_expl.font.size = Pt(11)
    return p

def add_code_block(code_text):
    """添加代码块（用灰色背景段落模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ============================================================
#                           封面
# ============================================================

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Qwen-VL-RS 项目完全指南')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于 Qwen3-VL 的遥感图像描述与理解\n—— 面向初学者的架构解析与名词详解')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('武汉大学电子信息学院 · 研0实习项目\n2026年6月\n文档版本 v1.0')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
#                    目录（手动）
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    '第一章  项目概览 —— 一句话讲清楚我们在做什么',
    '第二章  背景与动机 —— 为什么遥感需要自己的 VLM',
    '第三章  项目架构详解 —— 每个文件是干什么的',
    '第四章  核心概念名词解释（30+ 个术语，从零讲起）',
    '第五章  技术栈深度解读 —— 为什么选这些工具',
    '第六章  实验设计解读 —— 7 组消融实验分别回答什么问题',
    '第七章  上手路线图 —— 6 周从零到完整项目',
    '第八章  面试准备清单 —— 做完项目后面试官会问什么',
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# ============================================================
#              第一章：项目概览
# ============================================================
doc.add_heading('第一章  项目概览', level=1)

add_para('1.1 一句话总结', bold=True)
add_para('本项目使用 Qwen3-VL（阿里的多模态大模型），通过 LoRA 高效微调，让它学会"看懂遥感图像并用自然语言描述出来"。最终目标是：用 2B 参数的小模型，在遥感图像描述任务上超越 GPT-4V 的零样本表现。')

add_para('1.2 核心流程', bold=True)
add_para('整个项目的核心流程可以用一句话概括：')
add_para('"下载遥感图像数据集 → 手写 PyTorch 训练管线 → 用 LoRA 微调 Qwen-VL → 在多个 benchmark 上评估 → 分析模型哪里犯错 → 导出 Demo 展示"')
add_para('下面这张图可以帮助理解（文字版）：', size=10)
add_code_block('遥感卫星图  →  Qwen3-VL (视觉编码器)  →  LLM (LoRA微调)  →  "图中有一片农田，'+
               '东北方向有一条高速公路，右下角是居民区..."')

add_para('1.3 你为什么需要认真读这个文档', bold=True)
add_para('如果你只是"跑通了一个代码"，面试官问三个问题就会露馅。这份文档的目标是让你理解项目中每一个术语、每一个设计决策、每一个文件的作用——这样你在面试中能讲出 15 分钟以上不重复、有深度的内容。')

doc.add_page_break()

# ============================================================
#              第二章：背景与动机
# ============================================================
doc.add_heading('第二章  背景与动机', level=1)

add_para('2.1 遥感图像是什么？', bold=True)
add_para('遥感（Remote Sensing）是指通过卫星、无人机、飞机等平台搭载的传感器，在不接触目标的情况下获取地表信息的技术。你每天用的 Google Earth / 高德地图的卫星视图，就是遥感图像。')
add_para('遥感图像和普通照片有几个关键区别：')
add_bullet('尺度差异巨大：一张图可能覆盖 100 平方公里，也可能只看一个操场')
add_bullet('没有固定朝向：卫星从不同轨道拍摄，不像手机拍照永远是"上北下南"')
add_bullet('光谱信息丰富：除了可见光（RGB），还有近红外、热红外等波段')
add_bullet('目标密集且相似：农田和草地看起来很像，道路和河道也容易混淆')
add_bullet('空间上下文重要：地物之间的关系（比如"农田被森林环绕"）是描述的关键')

add_para('2.2 为什么通用 VLM 在遥感上表现不好？', bold=True)
add_para('GPT-4V、Qwen-VL 这些通用模型是在互联网图片（自拍、风景、商品图）上训练的。当你给它们看遥感图像时：')
add_bullet('它们分不清"农田"和"草地"——训练数据里几乎没有航拍角度的农田')
add_bullet('它们不理解"位于东北方向"——通用模型的空间描述能力本身就很弱')
add_bullet('它们不认识遥感专有名词——"归一化植被指数(NDVI)"、"合成孔径雷达(SAR)"等')
add_bullet('它们看不懂俯视图——训练数据 99% 是水平视角的照片')
add_para('这给了我们机会：用少量遥感数据微调一个小模型，让它在遥感这个细分领域超越 GPT-4V。这就是本项目要做的事。')

add_para('2.3 为什么不直接用 GPT-4V 的 API？', bold=True)
add_bullet('成本：处理 10,000 张遥感图像，GPT-4V API 费用可能上千元，而本地模型几乎零成本')
add_bullet('隐私：遥感数据可能涉密（军事基地、敏感设施），不能上传到 OpenAI 服务器')
add_bullet('延迟：API 调用有网络延迟，本地推理可以做到 << 1 秒')
add_bullet('可定制：你可以针对特定场景（如灾害评估、城市规划）继续微调，API 做不到')
add_bullet('学术价值：微调本身就是一个研究方向，面试官更看重你"怎么微调的"而不是"调了什么 API"')

doc.add_page_break()

# ============================================================
#              第三章：项目架构详解
# ============================================================
doc.add_heading('第三章  项目架构详解', level=1)

add_para('本章逐一介绍项目中的每个文件和目录，解释它们的具体职责。读完这一章，你应该能清楚地告诉别人"这个项目是怎么组织的"。')

# --- 3.1 根目录 ---
add_para('3.1 根目录文件', bold=True, size=13)

add_term('README.md', '项目的主文档。包含项目动机、结构说明、技术栈、实验设计（7 组消融实验 + 4 个 baseline）、6 周时间规划、面试要点。这是你面试时给面试官看的第一张"名片"。')
add_term('requirements.txt', 'Python 依赖清单。用 pip install -r requirements.txt 一键安装。包含了 torch、transformers、peft（LoRA 实现）、gradio（Web Demo）、vllm（高性能推理）等。')
add_term('.gitignore', 'Git 版本控制的忽略规则。排除 __pycache__、模型权重文件（*.safetensors）、训练日志、IDE 配置等不该提交到仓库的文件。')

# --- 3.2 configs ---
add_para('3.2 configs/ —— 配置文件目录', bold=True, size=13)

add_term('configs/sft_config.yaml', '训练主配置文件。定义了模型路径、LoRA 超参数（rank=16, alpha=32）、训练参数（学习率=2e-5，batch_size=8，epoch=3）、优化器配置等。训练时一行命令 python -m training.trainer --config configs/sft_config.yaml 就启动。')
add_term('configs/lora_config.yaml', 'LoRA 消融实验参数矩阵。定义了消融实验中要测试的变量：rank ∈ {8,16,32,64}、target_modules 的三种组合（仅 q+v / 全部注意力层 / 所有线性层）、alpha 和 dropout 的取值。')
add_term('configs/data_config.yaml', '数据集配置。定义了三个数据集（RSICD 10921 张 / UCM 2100 张 / Sydney 613 张）的路径、图像预处理参数（CLIP 归一化均值/标准差）、数据增强策略（遥感专用：离散旋转 0/90/180/270°、光谱扰动）、训练/验证/测试划分比例。')

add_para('什么是 YAML？', bold=True, size=10)
add_para('YAML（读作"雅缪"）是一种人类易读的配置文件格式。它用缩进表示层级，用冒号分隔键和值。比 JSON 更简洁（不需要引号和花括号），比命令行参数更清晰（所有的配置一目了然）。深度学习项目中常用 YAML 来管理超参数。', size=10)

# --- 3.3 data ---
add_para('3.3 data/ —— 数据处理模块', bold=True, size=13)

add_term('data/dataset.py', '核心的数据集加载器。它做的事情：① 读取 JSONL 格式的标注文件（每行一个 JSON，包含图像路径和 5 句参考描述）；② 按 8:1:1 划分训练/验证/测试集；③ 在 __getitem__ 中返回一张图像 + 对应的 tokenized 描述。支持三个数据集（RSICD / UCM / Sydney）的统一加载。')
add_term('data/transforms.py', '图像增强管线。遥感图像的增强和自然图像不同——遥感图像没有固定朝向，所以用离散旋转（0°/90°/180°/270°）；用多尺度裁剪模拟不同卫星的拍摄分辨率；用亮度/对比度扰动模拟不同大气条件。基于 albumentations 库实现。')
add_term('data/collator.py', 'Data Collator（批次整理器）。这是多模态训练中最容易出"静默 bug"的地方。它负责把一个 batch 内的图像和文本整理成模型可以直接吃的张量格式。特别是：图像张量堆叠、文本序列的变长填充（用 pad_token_id 填充到 batch 内最大长度）、attention_mask 的正确构建。')
add_term('data/prompts.py', 'Prompt 模板库。定义了三种 prompt 风格：① 标准型（"请描述这张遥感图像"）；② 遥感专家型（注入领域 context，如"你是遥感分析专家，请从地物类型、空间布局、人造结构三方面分析"）；③ 思维链型（引导模型逐步推理）。这是消融实验 A4 的研究对象——不同 prompt 对模型输出质量有显著影响。')
add_term('data/raw/ 和 data/processed/', '原始数据和预处理后数据的存放目录。raw/ 存放下载的数据集原始文件，processed/（可选）存放预处理后的缓存。这两个目录通过 .gitkeep 文件在 Git 中保持空目录结构，但实际数据（几个 GB 的图片）不会提交到 Git。')

# --- 3.4 models ---
add_para('3.4 models/ —— 模型模块', bold=True, size=13)

add_term('models/qwen_vl_rs.py', '模型的封装类。它做的事情：① 加载 Qwen3-VL-2B 预训练模型（视觉编码器 + LLM）；② 在 LLM 部分注入 LoRA adapter（通过 PEFT 库）；③ 提供 generate() 方法用于推理；④ 提供 merge_and_save() 方法将 LoRA 权重合并回基座模型。')
add_para('模型架构详解：', bold=True, size=10)
add_code_block('视觉编码器 (ViT, 冻结)     →  提取图像特征\n    ↓\nMLP 投影层 (可选 LoRA)     →  将视觉特征映射到 LLM 的输入空间\n    ↓\n大语言模型 (Qwen3, LoRA)  →  根据图像特征 + 文本 prompt 生成描述\n    ↓\n文本输出  →  "图中有一片农田，东北方向有高速公路..."')
add_para('关键设计决策：视觉编码器冻结（遥感图像的低级特征与自然图像共享，不需要重训），只微调 LLM 部分（因为领域词汇和描述风格需要适配）。这是典型的"参数高效微调"策略。', size=10)

# --- 3.5 training ---
add_para('3.5 training/ —— 训练模块', bold=True, size=13)

add_term('training/trainer.py', '训练编排器。基于 HuggingFace Trainer（而不是 MS-SWIFT 那样的黑盒框架）。它负责：① 加载模型和数据集；② 配置 TrainingArguments（batch size、学习率、epoch 数等）；③ 注入自定义的 loss 函数和 metrics callback；④ 执行训练循环；⑤ 保存最优模型（基于验证集 CIDEr-D 得分）。')
add_para('为什么要手写 Trainer 而不是用 MS-SWIFT？', bold=True, size=10)
add_para('MS-SWIFT 帮你省了 50 行代码，但你完全不知道训练循环里发生了什么。面试官问"你的 data collator 怎么处理 image token 的？""你的 loss 加了哪些 trick？"——用 SWIFT 的人答不出来。手写 Trainer 让你理解每一步，面试时这 是最有说服力的"我懂"的证明。', size=10)

add_term('training/loss.py', '自定义损失函数。提供了三种损失函数：① 标准交叉熵（默认）；② Label Smoothing 交叉熵（让模型不那么"自信"，防止过拟合）；③ Focal Loss（给难分类的 token 更大权重，帮助模型学习稀有词汇如"植被指数"）。')
add_term('training/metrics.py', '评估指标计算。实现了 5 个标准图像描述指标（BLEU / METEOR / ROUGE-L / CIDEr-D / SPICE）加上 2 个遥感特定指标（CHAIR 幻觉检测 / 地物类别 F1）。')

# --- 3.6 evaluation ---
add_para('3.6 evaluation/ —— 评估模块', bold=True, size=13)

add_term('evaluation/eval.py', '评估引擎。加载训练好的模型，在多个数据集的测试集上生成描述，然后与 ground truth 参考描述对比计算所有指标。')
add_term('evaluation/benchmarks.py', '多模型对比运行器。自动运行 4 个 baseline（Qwen-VL zero-shot / BLIP-2 LoRA / GPT-4V API / 我们的模型）并在同一数据集上对比，生成 Markdown 格式的对比表格。')
add_term('evaluation/error_analysis.py', '错误分析器。把模型预测的错误按维度分类：① 地物类别误判；② 空间关系描述错误；③ 物体计数偏差；④ 幻觉（描述了图中不存在的东西）；⑤ 光谱描述错误。找到"最差的 20 个 case"供人工检查——这是你和面试官聊"模型局限性"时最好的素材。')

# --- 3.7 inference ---
add_para('3.7 inference/ —— 推理模块', bold=True, size=13)

add_term('inference/infer.py', '推理引擎。支持单张图像推理和批量推理。加载模型后，输入一张遥感图像，输出描述文本。')
add_term('inference/gradio_app.py', 'Gradio 交互 Demo。提供一个 Web UI 界面：左侧上传图像 + 调节参数（温度、最大 token 数），右侧显示模型输出。支持与 GPT-4V / BLIP-2 的并排对比。面试时你可以直接打开 localhost:7860 给面试官演示。')
add_term('inference/vllm_serve.py', 'vLLM 高性能推理服务。使用 PagedAttention 技术实现高效的 KV-cache 管理，支持连续批处理（continuous batching），适合生产环境部署。当面试官问"如果要上线你怎么做"，你可以提这个。')

# --- 3.8 其他 ---
add_para('3.8 其他目录', bold=True, size=13)

add_term('scripts/', '四个 bash 脚本：train_lora.sh（LoRA 微调启动）、train_sft.sh（全量微调启动）、eval.sh（评估启动）、inference.sh（推理/Demo 启动）。一键运行，无需记忆 Python 命令行参数。')
add_term('experiments/', '实验记录目录。exp_template.md 是每次实验的记录模板（假设、配置、结果、观察、结论），logs/ 存放训练日志和 TensorBoard 文件。养成每次实验都认真记录的习惯——这是你面试时被问"你试过哪些没成功的方法"时的证据。')
add_term('notebooks/', 'Jupyter Notebook 目录。计划 4 个 notebook：数据探索、baseline zero-shot 分析、错误案例分析、消融实验可视化。Notebook 的好处是可以边写代码边写分析文字，适合探索性工作和展示。')
add_term('tests/', '单元测试目录。3 个测试文件分别测试数据集加载、指标计算、模型前向传播。虽然工业界代码覆盖率通常要求 80%+，但这里只需要保证核心逻辑正确即可。')

doc.add_page_break()

# ============================================================
#              第四章：核心概念名词解释
# ============================================================
doc.add_heading('第四章  核心概念名词解释', level=1)

add_para('本章是整份文档最有价值的部分。我们对项目中出现的 30+ 个专业术语逐一解释，用通俗的语言讲清楚每个概念"是什么、为什么需要、和项目有什么关系"。')

# --- 基础概念 ---
doc.add_heading('4.1 基础概念', level=2)

add_term('大语言模型 (LLM, Large Language Model)',
         '一种基于 Transformer 架构的神经网络，通过在数十亿字的文本上训练，学会了"理解语言"和"生成文本"。ChatGPT、Qwen、LLaMA 都是 LLM。简单理解为：一个超级智能的"文字接龙"机器，你给它一句话，它预测下一个字应该是什么。')
add_para('和项目的关系：Qwen3-VL 的 "语言部分" 就是一个 LLM，负责根据图像特征生成描述文字。我们微调的正是这部分。', size=10, indent=0.8)

add_term('视觉语言模型 (VLM, Vision-Language Model)',
         '在 LLM 的基础上加了一个"眼睛"（视觉编码器）。它可以同时理解图像和文本输入，并生成文本输出。GPT-4V、Qwen-VL、LLaVA、BLIP-2 都是 VLM。')
add_para('和项目的关系：Qwen3-VL-2B-Instruct 就是我们使用的 VLM 基座模型。', size=10, indent=0.8)

add_term('多模态 (Multi-modal)',
         '"模态"指的是数据的类型。文本是一种模态，图像是另一种模态，语音、视频也是。多模态模型就是能同时处理多种数据类型的模型。"多模态"不一定是"好"的——多模态模型的训练和评估都比单模态复杂得多。')
add_para('和项目的关系：我们的任务本质上是"图像→文本"的多模态生成任务（Image Captioning）。', size=10, indent=0.8)

add_term('Transformer',
         '2017 年 Google 提出的神经网络架构，用"自注意力机制"（Self-Attention）替代了传统的 RNN/LSTM。几乎所有现代 LLM 和 VLM 都基于 Transformer。它的核心思想是：让序列中的每个位置都能直接"关注"到序列中所有其他位置，而不是像 RNN 那样一步一步地传递信息。')
add_para('和项目的关系：Qwen3-VL 的内部架构就是 Transformer。LoRA 也正是作用于 Transformer 的注意力层参数矩阵上。', size=10, indent=0.8)

# --- SFT ---
doc.add_heading('4.2 训练相关术语', level=2)

add_term('预训练 (Pre-training)',
         '在大规模无标注（或弱标注）数据上训练模型，让它学会通用的语言/视觉表示。比如 Qwen3-VL 在几十亿图文对上学到了"狗长什么样""红色是什么意思"。预训练是烧钱最多的阶段，只有大公司/大实验室做得起。')
add_para('和项目的关系：我们直接使用 Qwen 团队已经预训练好的模型，不需要（也没钱）自己预训练。', size=10, indent=0.8)

add_term('微调 (Fine-tuning)',
         '在预训练模型的基础上，用少量、高质量的领域数据继续训练，让模型适应特定任务。微调的成本远低于预训练（一张消费级显卡就行），但效果可能非常显著。')
add_para('和项目的关系：整个项目就是一个"微调"项目——用遥感数据把通用 VLM 变成遥感专家。', size=10, indent=0.8)

add_term('SFT (Supervised Fine-Tuning, 监督微调)',
         '最基础的微调方式：给模型看"输入→期望输出"的配对数据，让它模仿学习。比如输入一张遥感图 + "请描述"，期望输出是人工写的描述。SFT 是后训练（Post-training）的第一步，也是最基础的一步。')
add_para('和项目的关系：我们的核心训练就是 SFT——给 Qwen-VL 看遥感图和对应的描述，让它学会"遥感描述"这个技能。', size=10, indent=0.8)

add_term('LoRA (Low-Rank Adaptation, 低秩适配)',
         '一种"参数高效微调"方法。它不修改原始模型权重，而是在旁边插入小的可训练矩阵（adapter）。原始模型权重被冻结（不动），只训练这些小矩阵。LoRA 的核心数学直觉：一个大矩阵的变化量 ΔW 可以用两个小矩阵 A×B 近似表示，需要训练的参数从 d×d 降到 d×(r+r)，通常 r ∈ {8, 16, 32}。')
add_para('和项目的关系：我们使用 LoRA 微调 Qwen-VL。原始模型 ~2B 参数，LoRA 只需要训练 ~1%（约 20M 参数），一张 12GB 显存的显卡就能跑。', size=10, indent=0.8)

add_term('QLoRA (Quantized LoRA)',
         'LoRA 的升级版：在 LoRA 的基础上，把原始模型量化到 4-bit 精度，进一步降低显存占用。适合显存特别紧张的情况。')
add_para('和项目的关系：目前我们用标准 LoRA（BF16 精度），如果显存不够可以考虑切换到 QLoRA。', size=10, indent=0.8)

add_term('Rank (秩)',
         '在线性代数中，矩阵的秩表示它的"信息丰富程度"。在 LoRA 中，rank 指定了 adapter 矩阵的维度。rank 越高 → adapter 容量越大 → 能学到更复杂的模式 → 但训练更慢、可能过拟合。rank 越低则反之。rank 是 LoRA 最重要的超参数，也是我们消融实验 A1 的研究对象。')
add_para('和项目的关系：默认 rank=16，消融实验中测试 {8, 16, 32, 64}。', size=10, indent=0.8)

add_term('DPO (Direct Preference Optimization, 直接偏好优化)',
         '一种对齐方法：不是告诉模型"正确答案是什么"（像 SFT 那样），而是告诉模型"A 比 B 更好"（偏好数据）。DPO 直接优化模型使其偏好 chosen（被偏好的回答）而非 rejected（被拒绝的回答）。常用于减少模型的啰嗦、改善格式遵循、消除不良输出。')
add_para('和项目的关系：本项目目前只做 SFT，但你的原始项目中做过 DPO。在面试时你可以提"后续可以加入 DPO 进一步减少遥感描述中的幻觉"。', size=10, indent=0.8)

add_term('RLHF (Reinforcement Learning from Human Feedback, 基于人类反馈的强化学习)',
         'ChatGPT 成功的核心技术之一。流程：① 收集人类偏好数据；② 训练一个 Reward Model 来预测人类会给多少"奖励分"；③ 用强化学习（PPO 算法）优化模型让 Reward Model 给的分数最大。RLHF 比 DPO 更复杂但效果可能更好。DPO 可以看作是 RLHF 的简化版。')
add_para('和项目的关系：本项目的"进阶扩展方向"——如果能收集到遥感专家的偏好标注，做一轮 RLHF/DPO 会显著提升描述的准确性和专业性。', size=10, indent=0.8)

add_term('Epoch (训练轮次)',
         '模型把整个训练数据集完整"看"一遍，叫做一个 epoch。通常需要多个 epoch 才能学好（本项目中设为 3 个 epoch）。epoch 太少→欠拟合；epoch 太多→过拟合。')
add_para('和项目的关系：默认 3 epoch，这个值需要根据验证集 loss 曲线来调整。', size=10, indent=0.8)

add_term('Batch Size (批次大小)',
         '每次更新模型参数时使用的样本数量。batch size 太小 → 训练不稳定；batch size 太大 → 显存不够。通常用"梯度累积"(Gradient Accumulation) 来模拟更大的 batch size。')
add_para('和项目的关系：本项目 per_device_batch_size=2，gradient_accumulation_steps=4，等效 batch size = 2×4 = 8。', size=10, indent=0.8)

add_term('学习率 (Learning Rate)',
         '控制模型参数每次更新的步长。太大 → 训练震荡不收敛；太小 → 训练太慢。学习率是训练最重要的超参数，通常用 warmup（前几百步从 0 线性增长到目标值）来稳定训练。')
add_para('和项目的关系：默认 2e-5（0.00002），配合 cosine 调度器（学习率按余弦曲线逐渐衰减）。', size=10, indent=0.8)

add_term('过拟合 (Overfitting)',
         '模型在训练集上表现很好，但在测试集（没见过的数据）上表现很差。就像背答案的学生，题目变一下就不会了。通常通过数据增强、dropout、early stopping 等手段来缓解。')
add_para('和项目的关系：我们通过 cross-dataset evaluation（在 A 数据集上训练，在 B 数据集上测试）来检验是否过拟合。如果跨数据集分数严重下降，说明过拟合了。', size=10, indent=0.8)

add_term('梯度检查点 (Gradient Checkpointing)',
         '一种"用时间换空间"的技术：训练时不保存所有中间激活值，需要时重新计算。显存占用大幅降低，但训练速度变慢约 20%。是消费级显卡训练大模型的必备技巧。')
add_para('和项目的关系：配置中 gradient_checkpointing=True，因为 Qwen3-VL 2B 模型在 12GB 显卡上不开这个跑不动。', size=10, indent=0.8)

add_term('Data Collator (数据整理器)',
         '深度学习训练流程中的一个组件，负责将一个 batch 的原始样本（图像 + 文本）整理成模型可以直接前向传播的张量格式。对多模态模型来说，collator 特别容易出错——如果图像和文本的 attention mask 没有正确对齐，训练可能看起来正常但推理结果完全错误。')
add_para('和项目的关系：data/collator.py 是本项目的关键基础设施之一。', size=10, indent=0.8)

# --- 评估 ---
doc.add_heading('4.3 评估相关术语', level=2)

add_term('Benchmark (基准测试)',
         '一组标准化的测试集和评估指标，用于客观比较不同模型的性能。Benchmark 的作用就像高考——不同学校的学生（模型）用同一张卷子（benchmark）来比较。没有 benchmark 的话，每个人都可以说自己"效果好"，没有可比性。')
add_para('和项目的关系：我们用 RSICD + UCM + Sydney 三个数据集作为 benchmark，和 Qwen-VL zero-shot、BLIP-2、GPT-4V 做公平对比。', size=10, indent=0.8)

add_term('Zero-shot (零样本)',
         '模型"从来没有见过这个任务的训练样本"，直接做推理。GPT-4V 做遥感描述就是 zero-shot——它没有专门学过遥感描述，但因为它见多识广，也能做个大概。Zero-shot 是微调的 baseline——如果微调后还不如 zero-shot，那训练就是失败的。')
add_para('和项目的关系：Qwen3-VL zero-shot 和 GPT-4V zero-shot 都是我们的 baseline。', size=10, indent=0.8)

add_term('图像描述 (Image Captioning)',
         'CV+NLP 交叉任务：输入一张图像，输出一句/段自然语言描述。图像描述比图像分类难得多——分类只需要输出一个标签，描述需要理解场景、对象、属性、关系并用流畅的语言表达出来。')
add_para('和项目的关系：这是本项目要解决的核心任务。', size=10, indent=0.8)

add_term('Ground Truth (GT, 真实标注)',
         '人工标注的"正确答案"。在图像描述任务中，每张图像通常有 5 句人工写的描述作为 GT。评估时把模型生成的描述和 GT 比较，看有多接近。')
add_para('和项目的关系：我们的训练数据中每张图有 5 句人工描述，评估时和这 5 句对比计算得分。', size=10, indent=0.8)

add_term('BLEU (Bilingual Evaluation Understudy)',
         '最经典的文本生成评估指标，最初用于机器翻译。计算方式：统计生成文本中有多少个 n-gram（连续 n 个词）出现在参考文本中，再用一个 brevity penalty（长度惩罚）防止模型只输出很短的高频词。BLEU-4 表示考虑 1-gram 到 4-gram 的组合。')
add_para('局限性：BLEU 只看词面匹配，不理解语义。比如"一条狗在跑"和"一只犬在奔跑"在 BLEU 看来完全不同。因此图像描述领域已不再以 BLEU 为首要指标。', size=10, indent=0.8)
add_para('和项目的关系：我们仍然计算 BLEU 作为参考指标，但首选的指标是 CIDEr-D。', size=10, indent=0.8)

add_term('CIDEr / CIDEr-D (Consensus-based Image Description Evaluation)',
         '专门为图像描述设计的评估指标，是目前这个领域最重要的指标。核心思想：① 用 TF-IDF 给每个 n-gram 加权——频繁出现在所有描述中的常见词（如"the""a"）权重低，特定于某张图的描述词权重高；② 把模型生成的 n-gram 向量和所有参考描述的共识向量做余弦相似度。CIDEr-D 是 CIDEr 的去重改进版。')
add_para('为什么 CIDEr 比 BLEU 好？比如一张机场图，所有参考描述都提到了"runway"（跑道）。如果模型也说"runway"，CIDEr 会给高分（因为"runway"对这张图有高区分度）；如果模型只说"plane"而没说"runway"，BLEU 可能看不出来差异，但 CIDEr 会惩罚。', size=10, indent=0.8)
add_para('和项目的关系：CIDEr-D 是我们训练时选择最优模型（best model selection）和早停（early stopping）的依据。', size=10, indent=0.8)

add_term('METEOR (Metric for Evaluation of Translation with Explicit ORdering)',
         'BLEU 的改进版。除了精确匹配，还考虑了同义词匹配（如"dog"和"canine"）、词干匹配（如"running"和"run"）、词序匹配。对图像描述评估通常比 BLEU 更合理。')

add_term('ROUGE (Recall-Oriented Understudy for Gisting Evaluation)',
         '最初为自动摘要设计的指标，侧重召回率（参考文本中多少内容被生成文本覆盖了）。ROUGE-L 基于最长公共子序列（LCS），对句子结构的相似性敏感。')

add_term('SPICE (Semantic Propositional Image Caption Evaluation)',
         '最"智能"的图像描述指标。它把描述文本解析成场景图（scene graph）——对象、属性、关系的三元组，然后在场景图级别比较生成文本和参考文本的相似度。SPICE 和人类判断的相关性通常最高，但计算也最复杂。')

add_term('CHAIR (Caption Hallucination Assessment with Image Relevance)',
         '专门评估"模型有没有在描述中编造不存在的物体"的指标。CHAIR-s：多少比例的句子包含幻觉物体。CHAIR-i：所有提到的物体中幻觉物体占多少。这是遥感描述中非常重要的指标——模型不能瞎说"有座桥"但实际上没有。')
add_para('和项目的关系：我们的 error_analysis.py 实现了 CHAIR 指标，因为在遥感领域，幻觉是一个严重问题。', size=10, indent=0.8)

add_term('消融实验 (Ablation Study)',
         '研究方法的经典范式：保持其他条件不变，逐个移除/改变某个组件，观察性能变化。目的是回答："这个组件到底有没有用？有多大用？"。消融实验是审稿人和面试官最看重的实验类型——它证明了你的设计决策有实证依据，不是拍脑袋决定的。')
add_para('和项目的关系：本项目设计了 7 组消融实验（A1-A7），分别研究 LoRA rank、目标模块、数据量、prompt 设计、数据增强、多数据集联合训练、LLM 解冻策略的影响。', size=10, indent=0.8)

add_term('Baseline (基线)',
         '用于对比的"基准方法"。如果你的方法比所有 baseline 都差，那你的方法没有价值。如果比 baseline 好，需要说清楚"好在哪里、好多少"。')
add_para('和项目的关系：本项目的 4 个 baseline：Qwen-VL zero-shot、GPT-4V zero-shot、BLIP-2 LoRA、以及我们的方法本身（作为后续改进的 baseline）。', size=10, indent=0.8)

add_term('幻觉 (Hallucination)',
         '大模型的经典问题：模型"信心满满地说出错误信息"。在 VLM 中，幻觉可能表现为描述了图中不存在的物体（"图中有一座桥"但实际没有），或者错误描述了物体之间的关系。幻觉是 VLM 部署到实际场景（如医疗、遥感、自动驾驶）的最大障碍之一。')
add_para('和项目的关系：我们用 CHAIR 指标量化幻觉，并在 error_analysis.py 中专门分析幻觉案例。', size=10, indent=0.8)

# --- 工具/框架 ---
doc.add_heading('4.4 工具与框架', level=2)

add_term('HuggingFace (HF, 抱抱脸)',
         '目前最流行的 AI 模型开源平台。提供：① Model Hub（数万个预训练模型，免费下载使用）；② Transformers 库（统一的模型加载/训练/推理接口）；③ Datasets 库（标准化的数据集加载）；④ PEFT 库（LoRA 等高效微调方法的实现）。')
add_para('和项目的关系：我们的数据处理（datasets）、模型加载（transformers）、LoRA 实现（PEFT）、训练循环（Trainer）全部基于 HuggingFace 生态。', size=10, indent=0.8)

add_term('PEFT (Parameter-Efficient Fine-Tuning)',
         'HuggingFace 的一个库，统一实现了 LoRA、Prefix Tuning、Prompt Tuning、IA3 等参数高效微调方法。')
add_para('和项目的关系：我们用 PEFT 的 LoraConfig 和 get_peft_model 来注入 LoRA adapter。', size=10, indent=0.8)

add_term('vLLM',
         'UC Berkeley 开发的高性能 LLM 推理框架。核心创新是 PagedAttention——把 KV-cache 像操作系统管理内存一样分页管理，大幅提高显存利用率和吞吐量。vLLM 是目前 LLM 生产部署的事实标准之一。')
add_para('和项目的关系：inference/vllm_serve.py 定义了 vLLM 服务接口（但实际没有实现——这是面试中展示工程思维的"钩子"）。', size=10, indent=0.8)

add_term('Gradio',
         '快速构建机器学习 Demo 的 Python 库。几行代码就能创建一个 Web UI，支持图像上传、文本输入、参数调节。面试时打开 localhost:7860 给面试官看效果，比贴截图有说服力得多。')
add_para('和项目的关系：inference/gradio_app.py 定义了 Demo 的结构，实际实现后可以一键启动。', size=10, indent=0.8)

add_term('TensorBoard',
         'Google 开发的可视化工具。在训练过程中实时显示 loss 曲线、学习率变化、指标变化等。你不需要一直盯着终端，打开浏览器就能看到训练状态。')
add_para('和项目的关系：训练配置中 report_to=["tensorboard"]，训练日志自动输出到 experiments/logs/。', size=10, indent=0.8)

add_term('DeepSpeed ZeRO',
         '微软开发的分布式训练优化框架。ZeRO（Zero Redundancy Optimizer）分三个阶段：ZeRO-1 切分优化器状态；ZeRO-2 切分优化器状态+梯度；ZeRO-3 切分优化器状态+梯度+参数。简单说就是让多张显卡协作训练一个大模型，每张卡只保存一部分参数。')
add_para('和项目的关系：单卡训练用不到（2B 模型 12GB 就够了），但面试时提一下"如果需要扩展到更大的模型（7B/72B），可以用 DeepSpeed ZeRO-3"是加分项。', size=10, indent=0.8)

add_term('MS-SWIFT',
         '阿里推出的模型训练框架，封装了 SFT、DPO、RLHF 等流程。优点是几行代码就能跑训练，缺点是像"黑盒"——你不知道里面发生了什么。')
add_para('注意：你的原始项目（D:\\Qwen）使用了 MS-SWIFT。但在本项目中我们故意不用，而是手写训练循环——这会让你在面试中更有优势。', size=10, indent=0.8)

add_term('Flash Attention',
         '一种高效的注意力计算算法。传统 self-attention 的显存和时间复杂度都是 O(n²)，Flash Attention 通过分块计算和重计算策略，在不牺牲精度的前提下大幅降低显存占用和加速训练。')
add_para('和项目的关系：我们的配置中 attn_implementation="sdpa"（PyTorch 原生的 scaled dot-product attention），如果 GPU 支持可以换成 "flash_attention_2" 以获得更好的性能。', size=10, indent=0.8)

# --- 数据 ---
doc.add_heading('4.5 数据相关术语', level=2)

add_term('RSICD (Remote Sensing Image Captioning Dataset)',
         '中科院发布的遥感图像描述数据集。包含 10,921 张遥感图像，每张有 5 句英文描述，涵盖 30+ 个场景类别（机场、海滩、桥梁、森林、居民区等）。是目前遥感描述领域最大的公开数据集。')
add_para('和项目的关系：RSICD 是我们的主力训练集（数据量最大），也是评估的首选 benchmark。', size=10, indent=0.8)

add_term('UCM-Captions',
         '基于 UC Merced Land Use Dataset（21 类土地利用图像）扩展的描述数据集。2,100 张 256×256 图像，每张 5 句描述。')
add_para('和项目的关系：UCM 是我们第二大的训练集，也是跨数据集泛化评估（cross-dataset evaluation）的重要测试集。', size=10, indent=0.8)

add_term('Sydney-Captions',
         '基于 Sydney Dataset（7 类悉尼遥感图像）扩展的描述数据集。613 张 500×500 图像，每张 5 句描述。')
add_para('和项目的关系：Sydney 是最小的数据集（613 张），主要用于"小样本微调效果"分析和跨数据集泛化测试。', size=10, indent=0.8)

add_term('JSONL (JSON Lines)',
         '每行一个独立 JSON 对象的文件格式。比单个大 JSON 文件更适合流式读写和逐行处理。深度学习中的训练数据通常都用 JSONL 格式存储。')
add_para('和项目的关系：三个数据集的标注文件都是 JSONL 格式。你的原始项目（D:\\Qwen）中的 SFT_data_fixed.jsonl 和 dpo_dataset_swift_final.jsonl 也是 JSONL。', size=10, indent=0.8)

add_term('数据增强 (Data Augmentation)',
         '通过对原始数据施加随机变换（旋转、裁剪、颜色扰动等）来"创造"更多训练数据。核心假设：这些变换不改变图像的语义标签（一张遥感图旋转 90° 后仍然是同一张遥感图）。数据增强是最有效的防过拟合手段之一。')
add_para('和项目的关系：我们设计了遥感专用的数据增强策略——离散旋转（非连续旋转，因为遥感图无固定朝向）、多尺度裁剪、光谱扰动。详见 configs/data_config.yaml。', size=10, indent=0.8)

add_term('Prompt / 指令模板',
         '"提示词"或"指令模板"。你给 VLM 的文本指令。Prompt 设计对 VLM 的输出质量有巨大影响——同一张图，"请描述"和"你是一个遥感专家，请分析..."得到的回答质量可能天差地别。')
add_para('和项目的关系：data/prompts.py 定义了三类 prompt（标准/专家/思维链），消融实验 A4 会对比它们的效果。', size=10, indent=0.8)

# --- 计算机视觉 ---
doc.add_heading('4.6 计算机视觉相关术语', level=2)

add_term('视觉编码器 (Vision Encoder / ViT)',
         'VLM 的"眼睛"部分。通常是一个 Vision Transformer（ViT），负责把一张图像转换成一个特征向量序列。这个过程叫做"视觉编码"。ViT 的原理：把图像切成固定大小的小块（patches），每个 patch 像 NLP 中的一个 token 一样送入 Transformer。')
add_para('和项目的关系：Qwen3-VL 的视觉编码器是冻结的（不做训练），因为我们假设遥感图像的低级视觉特征和自然图像是共享的。', size=10, indent=0.8)

add_term('Ground Sample Distance (GSD, 地面采样距离)',
         '遥感术语：相邻像素中心在地面上对应的实际距离。GSD 越小 → 分辨率越高。比如 0.5m GSD 意味着一个像素代表地面上 0.5×0.5 米的区域。不同卫星的 GSD 差异很大（从亚米到几十米）。')
add_para('和项目的关系：我们的多尺度裁剪增强就是为了让模型适应不同 GSD 的遥感图像。', size=10, indent=0.8)

add_term('NDVI (归一化植被指数)',
         '遥感领域最常用的植被指数。利用红光和近红外波段的反射率差异来计算：(NIR - Red) / (NIR + Red)。NDVI 越高，植被越茂盛。')
add_para('和项目的关系：这是遥感专有名词的典型例子——通用 VLM 不知道什么是 NDVI，微调后应该能学会。', size=10, indent=0.8)

doc.add_page_break()

# ============================================================
#              第五章：技术栈深度解读
# ============================================================
doc.add_heading('第五章  技术栈深度解读', level=1)

add_para('本章解释我们"为什么选择这些技术而不是其他替代方案"。面试官经常问这类问题——他们想看到你"有技术选型的判断力"，而不是"别人用什么我就用什么"。')

add_para('5.1 为什么选 Qwen3-VL 而不是其他 VLM？', bold=True, size=12)
add_bullet('轻量：2B 参数版本可以在单张 12GB 消费级显卡上训练，不需要租 GPU 云服务器')
add_bullet('开源可商用：Apache 2.0 协议，不像 LLaMA 有使用限制')
add_bullet('中英双语：对中文遥感描述（如"农田"、"居民区"）支持好')
add_bullet('架构现代：使用 Qwen3 的 GQA（Grouped Query Attention）等技术，推理效率高')
add_bullet('社区活跃：阿里魔搭社区有完善的文档和 issue 响应')

add_para('5.2 为什么选 LoRA 而不是 Full Fine-tune？', bold=True, size=12)
add_para('核心原因是"性价比"。Full fine-tune（全量微调）需要更新所有 2B 参数，至少需要 4×A100（约 40 万元硬件）。LoRA 只需要更新 ~1% 参数，一张 RTX 3060（约 2000 元）就能跑。')
add_para('但更重要的面试答案是：LoRA 让你可以同时训练多个不同配置的模型（不同 rank、不同数据配比），然后对比效果——这就是消融实验的基础。如果是 Full fine-tune，每个实验跑 3 天，6 周只能做 2-3 组实验，形不成完整的消融结论。')

add_para('5.3 为什么手写 Trainer 而不是用 MS-SWIFT？', bold=True, size=12)
add_para('最简单的回答：面试官会问"你的训练循环里 loss function 是怎么实现的？data collator 怎么处理 image token 和 text token 的对齐？"。用 MS-SWIFT 的 black box，你答不出来。手写的 Trainer，每一行代码都是你写的，你自然能讲清楚。')

add_para('5.4 为什么选 CIDEr-D 作为首要评估指标而不是 BLEU？', bold=True, size=12)
add_para('BLEU 只看 n-gram 的精确匹配，对语义差异不敏感。CIDEr-D 用 TF-IDF 加权 n-gram——频繁出现在所有描述中的无意义词（the, a, is）权重低，对特定图像有区分度的内容词（runway, forest, residential）权重高。这使 CIDEr-D 和人类对描述质量的判断具有最高的相关性。')
add_para('面试时说这句话就够了："BLEU 更适合机器翻译，CIDEr 是为图像描述设计的——它考虑了跨图像的 n-gram 区分度，和人类判断的相关系数更高。"')

doc.add_page_break()

# ============================================================
#              第六章：实验设计解读
# ============================================================
doc.add_heading('第六章  实验设计解读', level=1)

add_para('实验设计是项目含金量的核心体现。一个好的实验设计能回答"每个设计决策对最终效果有多大贡献"这个问题。以下是本项目设计的七组消融实验——每一组回答一个特定的问题。')

add_para('6.1 基线对比（Baseline Comparison）', bold=True, size=12)
add_para('我们要证明"微调 Qwen3-VL 对遥感描述任务是有效的"。和谁比呢？')
add_bullet('Qwen3-VL zero-shot：不训练直接用，证明"微调有没有用"')
add_bullet('GPT-4V zero-shot：商业最强 VLM，证明"小模型微调后能否逼近大模型"')
add_bullet('BLIP-2 LoRA：另一个 VLM 框架，证明"用 Qwen 是不是比用 BLIP-2 好"')

add_para('6.2 七组消融实验分别回答什么问题', bold=True, size=12)

add_para('A1 — LoRA rank 的影响', bold=True)
add_para('问题：rank 设多大最好？太小学不会，太大浪费算力还可能过拟合。', size=10)
add_para('测试：r ∈ {8, 16, 32, 64}，固定其他参数不变。', size=10)
add_para('预期：随着 rank 增加，效果先升后稳——在某一个 rank 后边际收益趋近于零。这个"拐点"就是最优 rank。', size=10)

add_para('A2 — LoRA 目标模块的影响', bold=True)
add_para('问题：应该在哪些层注入 LoRA？只注入注意力层还是所有线性层？', size=10)
add_para('测试：q+v only / q+k+v+o / all-linear 三种配置。', size=10)
add_para('预期：all-linear 通常效果最好（因为语言模型的 MLP 层也存储了大量知识），但训练参数更多。', size=10)

add_para('A3 — 训练数据量的影响', bold=True)
add_para('问题：数据量对效果的影响有多大？有没有"数据效率"的瓶颈？', size=10)
add_para('测试：从全部数据中随机采样 {25%, 50%, 75%, 100%} 分别训练。', size=10)
add_para('预期：效果随数据量增长但边际递减。这个实验能说明"要不要去收集更多数据"——如果 75%→100% 几乎没有提升，说明现有数据已经够了。', size=10)

add_para('A4 — Prompt 设计的影响', bold=True)
add_para('问题：不同 prompt 对模型输出质量的差异有多大？遥感专家 prompt 真的比通用 prompt 好吗？', size=10)
add_para('测试：标准 prompt vs 遥感专家 prompt vs 思维链 prompt。', size=10)
add_para('预期：专家 prompt > 思维链 > 标准 prompt。这个结果可以写在论文里作为"prompt engineering matters in domain-specific VLM"的证据。', size=10)

add_para('A5 — 数据增强的影响', bold=True)
add_para('问题：遥感特定的数据增强（离散旋转、光谱扰动）到底有没有用？', size=10)
add_para('测试：无增强 / 仅空间增强 / 空间+光谱 / 全部增强。', size=10)
add_para('预期：数据增强应该能带来小幅但一致的提升，特别是在"跨数据集泛化"场景下。', size=10)

add_para('A6 — 多数据集联合训练的影响', bold=True)
add_para('问题：多个遥感数据集一起训练，是"互相干扰"还是"互相增强"？', size=10)
add_para('测试：单数据集 vs UCM+Sydney vs UCM+Sydney+RSICD。', size=10)
add_para('预期：多数据集联合训练应该能提升泛化能力——不同数据集的标注风格差异实际上是一种正则化。', size=10)

add_para('A7 — LLM 解冻策略的影响', bold=True)
add_para('问题：除了 LoRA adapter，要不要把 LLM 本身的某些层也解冻训练？', size=10)
add_para('测试：LoRA only / LoRA + 解冻 LLM 最后 2 层 / 全量微调。', size=10)
add_para('预期：LoRA + 解冻最后几层可能比纯 LoRA 稍好，但训练显存和计算量显著增加——需要权衡性价比。', size=10)

doc.add_page_break()

# ============================================================
#              第七章：上手路线图
# ============================================================
doc.add_heading('第七章  上手路线图', level=1)

add_para('以下是一份 6 周的从零上手计划。每一周都有明确的目标和产出。')

doc.add_heading('第 1 周：环境搭建 + 数据集探索', level=2)
add_bullet('安装 Python 环境和所有依赖（requirements.txt）')
add_bullet('下载 RSICD / UCM / Sydney 数据集')
add_bullet('用 Jupyter Notebook 做数据探索：统计各类别分布、图像尺寸分布、描述长度分布')
add_bullet('加载 Qwen3-VL 模型，用 zero-shot 方式跑几张图看看效果（baseline 数据采集）')
add_para('产出：data_exploration.ipynb、baseline zero-shot 指标表', size=10)

doc.add_heading('第 2 周：搭建训练管线', level=2)
add_bullet('实现 data/dataset.py：JSONL 解析、train/val/test 划分、__getitem__')
add_bullet('实现 data/transforms.py：遥感图像增强管线')
add_bullet('实现 data/collator.py：多模态批次整理')
add_bullet('实现 training/trainer.py：基于 HF Trainer 的训练循环')
add_bullet('用一个小数据集（如 Sydney 613 张）做冒烟测试，确保能跑通一个 epoch')
add_para('产出：可运行的训练代码（MVP）', size=10)

doc.add_heading('第 3 周：LoRA 微调 + 消融实验 A1-A4', level=2)
add_bullet('在 RSICD 上跑完整的 LoRA 微调（3 epoch，约 2-4 小时）')
add_bullet('跑 LoRA rank 消融（r=8/16/32/64）')
add_bullet('跑 target modules 消融')
add_bullet('跑数据量消融')
add_bullet('跑 prompt 消融')
add_bullet('用 TensorBoard 跟踪所有实验的 loss 和指标')
add_para('产出：消融实验结果表、训练曲线图', size=10)

doc.add_heading('第 4 周：完整评估 + 错误分析', level=2)
add_bullet('实现 evaluation/eval.py：全数据集评估')
add_bullet('实现 evaluation/benchmarks.py：多模型对比')
add_bullet('实现 evaluation/error_analysis.py：错误分类 + 幻觉检测')
add_bullet('跑 GPT-4V API 做对比评估（如果预算允许）')
add_bullet('人工检查 20 个最差 case，总结模型的典型错误模式')
add_para('产出：评估报告、错误分类统计、GPT-4V 对比表', size=10)

doc.add_heading('第 5 周：消融 A5-A7 + 最优模型 + Demo', level=2)
add_bullet('跑数据增强消融（A5）')
add_bullet('跑多数据集联合训练消融（A6）')
add_bullet('跑 LLM 解冻策略消融（A7）')
add_bullet('根据所有消融结果确定最优配置')
add_bullet('实现 inference/gradio_app.py')
add_bullet('训练最终最优模型')
add_para('产出：完整消融结论、最优模型、可运行的 Gradio Demo', size=10)

doc.add_heading('第 6 周：整理 + 文档 + 面试准备', level=2)
add_bullet('整理所有实验数据，生成最终的结论表格')
add_bullet('完善 GitHub README（中英文双语）')
add_bullet('准备 15 分钟的项目介绍 PPT / slides')
add_bullet('准备好面试中可能被追问的 20+ 个问题和答案')
add_bullet('如果有余力，写一篇技术博客（知乎/CSDN/博客园）')
add_para('产出：面试 ready 的完整项目', size=10)

doc.add_page_break()

# ============================================================
#              第八章：面试准备清单
# ============================================================
doc.add_heading('第八章  面试准备清单', level=1)

add_para('以下是面试官最可能问的问题——每一个都准备了参考答案框架。建议你把这些问题的答案用自己的话写下来，对着镜子/录音练习。')

doc.add_heading('8.1 项目介绍（必问）', level=2)
add_para('"请用 3 分钟介绍一下你的这个项目。"', bold=True)
add_para('框架：背景（遥感图像理解的挑战）→ 方法（Qwen3-VL + LoRA 微调）→ 实验（7 组消融 + 4 个 baseline）→ 结果（关键数字）→ 收获（最大的技术挑战是什么）。')

doc.add_heading('8.2 技术深度追问（高频）', level=2)

add_para('Q: "为什么不直接用 GPT-4V？"', bold=True)
add_para('A: 成本（API 调用费用）+ 隐私（遥感数据可能涉密）+ 延迟（本地推理 << 1s）+ 可定制性（可以针对特定场景继续微调）+ 学术价值（研究微调方法本身）。')

add_para('Q: "LoRA 的原理是什么？为什么它能工作？"', bold=True)
add_para('A: LoRA 假设模型在适应下游任务时，权重矩阵的变化量 ΔW 是低秩的，可以用两个小矩阵 A 和 B 的乘积近似。原始权重 W 冻结不动，只训练 A 和 B。这样可训练参数数量从 d×d 降到 d×(r+r)，r 通常取 8-64。')

add_para('Q: "你怎么确定模型没有过拟合？"', bold=True)
add_para('A: 三个方法：① 监控验证集 loss 和指标曲线（如果验证集指标开始下降但训练集还在涨，就是过拟合了）；② cross-dataset evaluation（在 RSICD 上训练，在 UCM 上测试，看泛化能力）；③ 数据增强消融实验（有增强 vs 无增强的对比可以说明正则化的效果）。')

add_para('Q: "你在这个项目中遇到的最大技术挑战是什么？"', bold=True)
add_para('A: 准备一个具体的故事。比如："多模态 Data Collator 的 attention mask 构建。最初我没注意 image token 占据的序列位置，导致 attention mask 和 input_ids 的对应关系错位——训练 loss 正常下降但推理结果全乱。后来通过逐行调试 attention mask 的每一维，发现问题出在 image token 的 padding 逻辑上。这个经历让我深刻理解了多模态模型中视觉 token 和文本 token 的序列构建机制。"')
add_para('这个故事的关键：① 有具体的技术问题；② 有调试过程；③ 有最终 solution；④ 有学到的教训。', size=10)

add_para('Q: "CIDEr-D 和 BLEU 的区别是什么？为什么选 CIDEr 作为首要指标？"', bold=True)
add_para('A: BLEU 是 n-gram 精确匹配 + 长度惩罚，不考虑词的区分度。CIDEr-D 用 TF-IDF 给 n-gram 加权——在所有描述中都出现的高频词权重低，特定于某张图的区分性词权重高。这使 CIDEr-D 和人类对图像描述质量的判断具有更高的相关性。')

add_para('Q: "如果要上线部署，还缺什么？"', bold=True)
add_para('A: ① 输入校验（非法图像、超大图像的处理）；② 响应式推理（用 vLLM 做 continuous batching）；③ 模型热更新（不停服切换新模型）；④ A/B 测试框架（对比新旧模型在生产环境的表现）；⑤ 监控告警（推理延迟、GPU 利用率、输出异常检测）；⑥ 版本管理（模型和代码的版本对应关系）。')

doc.add_heading('8.3 行为面试（中等频率）', level=2)

add_para('Q: "你为什么选择这个方向（遥感+VLM）？"', bold=True)
add_para('A: 结合你的背景——武大电子信息、遥感是武大王牌学科、你对多模态 AI 感兴趣。三者的交集就是这个项目。面试官想看到的是"你的项目选择有内在逻辑，不是随便选的"。"我是学电子信息的，武大遥感全国顶尖，我对大模型感兴趣，所以选了遥感+多模态这个交叉方向。"这个回答就够了。')

add_para('Q: "你试过哪些不成功的方法？"', bold=True)
add_para('A: 这是最容易暴露"你没有认真做实验"的问题。你必须准备至少 2 个具体的失败尝试。比如：① "我试过把 batch size 设成 1 然后不开 gradient checkpointing，结果 12GB 显存直接 OOM"；② "我试过用连续旋转（0-360°随机）做数据增强，但效果反而比离散旋转（0/90/180/270°）差——因为遥感图像中的地物方向和真实地理方位有关，连续旋转破坏了这种隐含的方向信息。"')

add_para('Q: "你有看过这个领域的最新论文吗？"', bold=True)
add_para('A: 准备 3 篇相关论文的名字和关键贡献。RSICD (2018)、RemoteCLIP (2024)、加上一篇你真正读过的（比如"我最近在看 GEO这篇文章，它做的端到端遥感目标检测..."）。不需要读过 20 篇——但需要真正读过 2-3 篇。')

doc.add_page_break()

# ============================================================
#              附录
# ============================================================
doc.add_heading('附录', level=1)

doc.add_heading('A. 常用命令速查', level=2)
add_code_block('# 安装依赖\npip install -r requirements.txt\n\n# LoRA 微调\nbash scripts/train_lora.sh\n\n# 评估\nbash scripts/eval.sh --checkpoint output/qwen_vl_rs_lora/best\n\n# Gradio Demo\npython inference/gradio_app.py\n\n# 查看 TensorBoard\ntensorboard --logdir experiments/logs')

doc.add_heading('B. 推荐学习资源', level=2)
add_bullet('《Attention Is All You Need》(2017) — Transformer 原始论文，必读')
add_bullet('LoRA 论文：Hu et al., "LoRA: Low-Rank Adaptation of Large Language Models" (2021)')
add_bullet('HuggingFace 官方教程：https://huggingface.co/learn/nlp-course')
add_bullet('Qwen-VL 官方文档：https://github.com/QwenLM/Qwen-VL')
add_bullet('RSICD 数据集论文：Lu et al., "RSICD: Remote Sensing Image Captioning Dataset" (2018)')
add_bullet('刘知远等《大模型十问》——了解 LLM 核心概念的最佳中文入门')

doc.add_heading('C. 文档更新记录', level=2)
add_para('v1.0 (2026-06): 初版，包含完整的项目架构解析和 30+ 术语解释。')

# ==================== 保存 ====================
output_path = 'D:/work/Qwen-VL-RS/Qwen-VL-RS-项目完全指南.docx'
doc.save(output_path)
print(f'[OK] 文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
